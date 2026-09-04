import gradio as gr
import pandas as pd
import os
import docx
import pdfplumber
import re
import zipfile
import time
import base64
import json
import urllib.parse
import urllib.request
from datetime import datetime
from groq import Groq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
from langchain_chroma import Chroma
from duckduckgo_search import DDGS
from gradio_client import Client
from huggingface_hub import InferenceClient

# ==========================================
# 1. CHAVES MESTRES E CONEXÕES
# ==========================================
chave_groq = os.environ.get("GROQ_API_KEY")
chave_hf = os.environ.get("HF_TOKEN")

cliente_groq = Groq(api_key=chave_groq)
cliente_hf = InferenceClient(token=chave_hf)
MODELO_GROQ = "llama-3.3-70b-versatile"
MODELO_VISAO = "llama-3.2-90b-vision-preview"

embeddings = HuggingFaceInferenceAPIEmbeddings(api_key=chave_hf, model_name="sentence-transformers/all-MiniLM-L6-v2")

# ==========================================
# 2. DIRETÓRIOS E BANCO DE DADOS LOCAL
# ==========================================
DIRETORIO = "./Central_IA_Master"
DIR_CHROMA = f"{DIRETORIO}/Banco_de_Dados_Vetorial"
DIR_CASOS = f"{DIRETORIO}/Projetos_Salvos"
DIR_MIDIA = f"{DIRETORIO}/Midia_Criada"
DIR_CHATS = f"{DIRETORIO}/Historico_Chats"

for d in [DIRETORIO, DIR_CASOS, DIR_MIDIA, DIR_CHATS]:
    os.makedirs(d, exist_ok=True)

# ==========================================
# 2.5. RADAR DA LOGOMARCA OFICIAL
# ==========================================
TAG_LOGO = ""
FAVICON_TAGS = ""
caminho_logo = None

for root_dir, _, files in os.walk("."):
    for f in files:
        if "chamariz" in f.lower() and f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')):
            caminho_logo = os.path.join(root_dir, f)
            break
    if caminho_logo:
        break

if caminho_logo:
    with open(caminho_logo, "rb") as f:
        b64_logo = base64.b64encode(f.read()).decode('utf-8')
        TAG_LOGO = f'<img src="data:image/png;base64,{b64_logo}" style="max-height: 85px; max-width: 100%; margin: 0 auto; display: block; filter: drop-shadow(0px 0px 15px rgba(212, 175, 55, 0.5));" alt="Código de Ouro" />'
        FAVICON_TAGS = f"""
        <link rel="icon" type="image/png" href="data:image/png;base64,{b64_logo}">
        <link rel="apple-touch-icon" href="data:image/png;base64,{b64_logo}">
        """
else:
    TAG_LOGO = '<div style="color:#D4AF37; text-align:center; font-weight:900; font-size: 24px; letter-spacing: 2px;">CÓDIGO DE OURO</div>'

# --- GESTÃO DE SESSÕES ---
def listar_sessoes_chat():
    sessoes = [f.replace('.json', '') for f in os.listdir(DIR_CHATS) if f.endswith('.json')]
    sessoes.sort(reverse=True)
    return sessoes if sessoes else ["Nenhuma conversa salva"]

def carregar_sessao_chat(id_sessao):
    if not id_sessao or id_sessao == "Nenhuma conversa salva": return [], id_sessao
    try:
        with open(f"{DIR_CHATS}/{id_sessao}.json", "r", encoding="utf-8") as f: return json.load(f), id_sessao
    except: return [], id_sessao

def iniciar_novo_chat():
    novo_id = f"Chat_{datetime.now().strftime('%d%m_%H%M%S')}"
    return [], novo_id, gr.update(choices=listar_sessoes_chat(), value=novo_id)

def listar_arquivos_mortos():
    arquivos = []
    for d in [DIR_CASOS, DIR_MIDIA]:
        if os.path.exists(d):
            for root, _, files in os.walk(d):
                for f in files:
                    if f.endswith(('.docx', '.pdf', '.mp3', '.jpg', '.mp4')): arquivos.append(os.path.join(root, f))
    arquivos.sort(key=os.path.getmtime, reverse=True)
    return arquivos

def atualizar_galeria_imagens():
    imgs = []
    for d in [DIR_CASOS, DIR_MIDIA]:
        if os.path.exists(d):
            for root, _, files in os.walk(d):
                for f in files:
                    if f.endswith(('.jpg', '.png', '.jpeg', '.webp')): imgs.append(os.path.join(root, f))
    imgs.sort(key=os.path.getmtime, reverse=True)
    return imgs

# ==========================================
# 3. EXTRAÇÃO MULTIMÍDIA E DADOS
# ==========================================
def encode_file_b64(caminho):
    with open(caminho, "rb") as f: return base64.b64encode(f.read()).decode('utf-8')

def extrair_texto(arquivo):
    caminho = arquivo.name if hasattr(arquivo, 'name') else arquivo
    texto = ""
    try:
        if caminho.lower().endswith('.pdf'):
            with pdfplumber.open(caminho) as pdf:
                for p in pdf.pages:
                    txt = p.extract_text()
                    if txt: texto += txt + "\n"
        elif caminho.lower().endswith(('.xlsx', '.csv')): texto = (pd.read_excel(caminho) if caminho.lower().endswith('.xlsx') else pd.read_csv(caminho)).to_string() 
        elif caminho.lower().endswith('.docx'):
            for p in docx.Document(caminho).paragraphs: texto += p.text + "\n"
        elif caminho.lower().endswith(('.mp3', '.ogg', '.wav')):
            with open(caminho, "rb") as file: texto = f"[TRANSCRIÇÃO]: {cliente_groq.audio.transcriptions.create(file=(caminho, file.read()), model='whisper-large-v3').text}\n"
        return texto
    except: return ""

# ==========================================
# 4. MOTORES DE MÍDIA
# ==========================================
def motor_gerar_imagem(prompt_desc, proporcao="Vertical"):
    try:
        w, h = (1080, 1920) if "vertical" in proporcao.lower() else (1920, 1080) if "horizontal" in proporcao.lower() else (1024, 1024)
        prompt_clean = urllib.parse.quote(prompt_desc.strip())
        url = f"https://image.pollinations.ai/prompt/{prompt_clean}?width={w}&height={h}&nologo=true&seed={int(time.time())}"
        nome_arq = f"{DIR_MIDIA}/Img_{datetime.now().strftime('%H%M%S')}.jpg"
        urllib.request.urlretrieve(url, nome_arq)
        return nome_arq
    except: return None

def motor_editar_imagem(caminho_imagem, prompt_edicao):
    try:
        res = cliente_hf.image_to_image(image=caminho_imagem, prompt=prompt_edicao, model="timbrooks/instruct-pix2pix")
        caminho_saida = f"{DIR_MIDIA}/Edit_{datetime.now().strftime('%H%M%S')}.jpg"
        res.save(caminho_saida)
        return caminho_saida
    except: return None

def motor_gerar_audio(texto):
    try:
        cam_txt = f"{DIR_MIDIA}/temp_{datetime.now().strftime('%H%M%S')}.txt"
        cam_audio = f"{DIR_MIDIA}/Voz_{datetime.now().strftime('%H%M%S')}.mp3"
        with open(cam_txt, "w", encoding="utf-8") as f: f.write(texto[:2500].replace('*', ''))
        os.system(f'edge-tts --voice pt-BR-AntonioNeural -f "{cam_txt}" --write-media "{cam_audio}"')
        return cam_audio
    except: return None

def motor_gerar_video(prompt_cena, imagem_base=None):
    try:
        if imagem_base: return Client("multimodalart/stable-video-diffusion", hf_token=chave_hf).predict(imagem_base, api_name="/video")
        else: return Client("multimodalart/zeroscope-v2", hf_token=chave_hf).predict(prompt_cena[:150], api_name="/infer")
    except: return None

# ==========================================
# 5. O CHAT AGÊNTICO
# ==========================================
def responder_chat_central(mensagem, historico, persona, usar_internet, id_sessao):
    texto_usuario = mensagem.get("text", "") if isinstance(mensagem, dict) else str(mensagem)
    arquivos = mensagem.get("files", []) if isinstance(mensagem, dict) else []
    contexto_extra = ""
    imagens_anexadas = []
    
    yield "⏳ *Sincronizando com a base de dados...*"
    
    for arq in arquivos:
        ext = arq.lower()
        if ext.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            imagens_anexadas.append(arq)
            yield f"👁️ *Analisando a imagem...*"
        else:
            yield f"📄 *Extraindo dados do documento...*"
            contexto_extra += f"\n[DOCUMENTO]:\n{extrair_texto(arq)}\n"
            
    if usar_internet and texto_usuario:
        yield "🌐 *Pesquisando na internet...*"
        try:
            resultados = DDGS().text(texto_usuario, max_results=4)
            contexto_extra += "\n\n[WEB]:\n" + "\n".join([f"{r['title']} - {r['body']}" for r in resultados])
        except: pass

    yield "🧠 *Processando IA...*"

    sys_prompt = f"""Você é a IA "Código de Ouro" operando no perfil {persona}.
Aja de forma inteligente, direta e com alto nível de execução.
PODERES EXECUTIVOS NO CHAT:
1. GERAR IMAGEM: [AÇÃO_IMAGEM: prompt em inglês detalhado 8k photorealistic | vertical]
2. EDITAR IMAGEM: [AÇÃO_EDITAR_IMAGEM: comando em inglês]
3. GERAR VÍDEO: [AÇÃO_VIDEO: prompt curto em inglês]
4. GERAR ÁUDIO: [AÇÃO_AUDIO: texto para falar em português]"""

    mensagens = [{"role": "system", "content": sys_prompt}]
    
    if historico:
        for item in historico:
            if isinstance(item, dict):
                mensagens.append({"role": item.get("role", "user"), "content": str(item.get("content", ""))})
            elif isinstance(item, (list, tuple)) and len(item) == 2:
                if item[0]: mensagens.append({"role": "user", "content": str(item[0])})
                if item[1]: mensagens.append({"role": "assistant", "content": str(item[1])})

    texto_final = (texto_usuario + contexto_extra).strip()

    if imagens_anexadas:
        conteudo_multimodal = [{"type": "text", "text": texto_final if texto_final else "Analise esta imagem."}]
        for img in imagens_anexadas:
            conteudo_multimodal.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encode_file_b64(img)}"}})
        mensagens.append({"role": "user", "content": conteudo_multimodal})
        modelo_escolhido = MODELO_VISAO
    else:
        mensagens.append({"role": "user", "content": texto_final if texto_final else "Olá!"})
        modelo_escolhido = MODELO_GROQ

    stream = cliente_groq.chat.completions.create(messages=mensagens, model=modelo_escolhido, max_tokens=4000, stream=True)
    
    resposta_acumulada = ""
    for pedaco in stream:
        delta = pedaco.choices[0].delta.content
        if delta:
            resposta_acumulada += delta
            yield re.sub(r'\[AÇÃO_\w+:.*?\]', '⚙️ *(Acionando motor multimídia...)*', resposta_acumulada)

    anexos_html = ""
    match_img = re.search(r'\[AÇÃO_IMAGEM:\s*(.*?)(?:\|\s*(\w+))?\]', resposta_acumulada)
    if match_img:
        prompt_i = match_img.group(1).strip()
        prop_i = "Vertical" if (match_img.group(2) and "vertical" in match_img.group(2).lower()) else "Quadrado"
        cam_gerada = motor_gerar_imagem(prompt_i, prop_i)
        if cam_gerada:
            b64_img = encode_file_b64(cam_gerada)
            anexos_html += f"\n\n**🖼️ Imagem:**\n<img src='data:image/jpeg;base64,{b64_img}' style='max-width:100%; border-radius:15px; border: 1px solid #D4AF37; margin-top:10px;' />\n"

    match_edit = re.search(r'\[AÇÃO_EDITAR_IMAGEM:\s*(.*?)\]', resposta_acumulada)
    if match_edit and imagens_anexadas:
        prompt_e = match_edit.group(1).strip()
        cam_edit = motor_editar_imagem(imagens_anexadas[-1], prompt_e)
        if cam_edit:
            b64_img = encode_file_b64(cam_edit)
            anexos_html += f"\n\n**✨ Edição:**\n<img src='data:image/jpeg;base64,{b64_img}' style='max-width:100%; border-radius:15px; border: 1px solid #D4AF37; margin-top:10px;' />\n"

    match_aud = re.search(r'\[AÇÃO_AUDIO:\s*(.*?)\]', resposta_acumulada)
    if match_aud:
        texto_loc = match_aud.group(1).strip()
        cam_aud = motor_gerar_audio(texto_loc)
        if cam_aud:
            b64_aud = encode_file_b64(cam_aud)
            anexos_html += f"\n\n**🔊 Áudio:**\n<audio controls src='data:audio/mp3;base64,{b64_aud}' style='width:100%; margin-top:10px;'></audio>\n"

    match_vid = re.search(r'\[AÇÃO_VIDEO:\s*(.*?)\]', resposta_acumulada)
    if match_vid:
        prompt_v = match_vid.group(1).strip()
        img_referencia = imagens_anexadas[-1] if imagens_anexadas else None
        cam_vid = motor_gerar_video(prompt_v, img_referencia)
        if cam_vid:
            b64_vid = encode_file_b64(cam_vid)
            anexos_html += f"\n\n**🎥 Vídeo:**\n<video controls style='max-width:100%; border-radius:15px; border: 1px solid #D4AF37; margin-top:10px;' src='data:video/mp4;base64,{b64_vid}'></video>\n"

    resposta_final_limpa = re.sub(r'\[AÇÃO_\w+:.*?\]', '', resposta_acumulada).strip() + anexos_html
    yield resposta_final_limpa

    try:
        sessao_alvo = id_sessao if (id_sessao and id_sessao != "Nenhuma conversa salva") else f"Chat_{datetime.now().strftime('%d%m_%H%M%S')}"
        arq_sessao = f"{DIR_CHATS}/{sessao_alvo}.json"
        historico_atual = []
        if os.path.exists(arq_sessao):
            with open(arq_sessao, "r", encoding="utf-8") as f: historico_atual = json.load(f)
        historico_atual.append({"role": "user", "content": texto_usuario if texto_usuario else "[Arquivo/Mídia]"})
        historico_atual.append({"role": "assistant", "content": resposta_final_limpa})
        with open(arq_sessao, "w", encoding="utf-8") as f: json.dump(historico_atual, f, ensure_ascii=False, indent=4)
    except: pass

def exportar_conversa_docx(historico):
    if not historico: return None
    pasta = f"{DIR_CASOS}/Exportacoes_{datetime.now().strftime('%d_%m_%H%M')}"
    os.makedirs(pasta, exist_ok=True)
    cam_word = f"{pasta}/Chat.docx"
    doc = docx.Document()
    doc.add_heading('Registro - Código de Ouro', 0)
    for item in historico:
        if isinstance(item, dict):
            autor = "Você:" if item.get("role") == "user" else "Código de Ouro:"
            doc.add_heading(autor, level=2)
            doc.add_paragraph(re.sub(r'<.*?>', '', item.get("content", "")))
    doc.save(cam_word)
    return cam_word

def gerar_dossie_lote(arquivos, instrucao, progresso=gr.Progress()):
    if not instrucao: return "⚠️ Forneça as instruções.", None, ""
    palavras = 0
    try:
        progresso(0.1, desc="Lendo documentos...")
        pasta = f"{DIR_CASOS}/Analise_{datetime.now().strftime('%d_%m_%Y__%Hh%M')}"
        os.makedirs(pasta, exist_ok=True)
        banco = Chroma(persist_directory=DIR_CHROMA, embedding_function=embeddings)
        fatiador = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
        
        if arquivos:
            for arq in arquivos:
                txt = extrair_texto(arq)
                palavras += len(txt.split())
                banco.add_texts([f"[FONTE: {os.path.basename(arq.name)}]\n{c}" for c in fatiador.split_text(txt)])
            
        progresso(0.5, desc="Analisando dados...")
        contexto = "\n".join([doc.page_content for doc in banco.similarity_search(instrucao, k=8)])
        prompt = f"DADOS:\n{contexto}\n\nINSTRUÇÃO: {instrucao}"
        resposta = cliente_groq.chat.completions.create(messages=[{"role": "user", "content": prompt}], model=MODELO_GROQ, max_tokens=4000).choices[0].message.content
        
        cam_word = f"{pasta}/Resultado.docx"
        doc = docx.Document()
        doc.add_heading('Análise Oficial', 0)
        doc.add_paragraph(resposta)
        doc.save(cam_word)
        return "✅ Concluído!", cam_word, resposta
    except Exception as e: return f"Erro: {e}", None, ""

# ==========================================
# 6. ARQUITETURA DE DESIGN TITAN (OBSIDIAN & GOLD)
# ==========================================
# Reset Base para Gradio
tema_titan = gr.themes.Base(font=[gr.themes.GoogleFont("Inter"), "sans-serif"]).set(
    body_background_fill="#050505", body_background_fill_dark="#050505",
    background_fill_primary="#050505", background_fill_primary_dark="#050505",
    background_fill_secondary="#0E0E12", background_fill_secondary_dark="#0E0E12",
    block_background_fill="#0E0E12", block_background_fill_dark="#0E0E12",
    border_color_primary="rgba(255, 255, 255, 0.05)", border_color_primary_dark="rgba(255, 255, 255, 0.05)",
    block_border_width="1px", block_radius="20px",
    body_text_color="#FFFFFF", body_text_color_dark="#FFFFFF",
    body_text_color_subdued="#888888", body_text_color_subdued_dark="#888888",
    color_accent_soft="rgba(212, 175, 55, 0.1)", color_accent_soft_dark="rgba(212, 175, 55, 0.1)",
)

PWA_HEAD = f"""
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
<meta name="theme-color" content="#050505">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-title" content="Código Ouro">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
{FAVICON_TAGS}
"""

LOGIN_HACK = """
<style>
    body, main, .wrap { background-color: #050505 !important; color: #fff !important; }
    form { background: #0E0E12 !important; border: 1px solid rgba(212,175,55,0.2) !important; border-radius: 24px !important; box-shadow: 0 20px 40px rgba(0,0,0,0.8) !important; padding: 45px !important; max-width: 90% !important; margin: auto !important; width: 420px;}
    button.primary { background: linear-gradient(135deg, #D4AF37, #AA7C11) !important; color: #000 !important; font-weight: 800 !important; border-radius: 12px !important; border: none !important; font-size: 16px !important; margin-top: 20px !important; transition: all 0.3s ease !important; width: 100%; padding: 14px !important;}
    button.primary:hover { transform: translateY(-2px); box-shadow: 0 8px 25px rgba(212,175,55,0.4) !important; }
    input { background-color: #050505 !important; border: 1px solid #222 !important; border-radius: 12px !important; color: #D4AF37 !important; padding: 14px !important; width: 100%; transition: border 0.3s;}
    input:focus { border-color: #D4AF37 !important; outline: none; }
    form h2 { display: none !important; }
</style>
<div style="text-align: center; margin-bottom: 30px; width: 100%;">
    [LOGO_PLACEHOLDER]
    <h1 style="color: #D4AF37; font-size: clamp(22px, 5vw, 28px); font-weight: 900; margin: 0; letter-spacing: 3px; font-family: 'Inter', sans-serif;">CÓDIGO DE OURO</h1>
    <p style="color: #666; font-size: 11px; margin-top: 8px; font-weight: 700; letter-spacing: 4px; text-transform: uppercase;">Acesso Restrito</p>
</div>
""".replace("[LOGO_PLACEHOLDER]", TAG_LOGO)

CSS_TITAN = """
/* RESET TOTAL PARA NÃO QUEBRAR O GRADIO */
body, html { font-family: 'Inter', sans-serif !important; background-color: #050505 !important; color: #F0F0F0 !important; }
footer { display: none !important; }

/* FIX CORES NATIVAS (Dropdowns, Inputs, Checkbox) */
input[type="checkbox"] { appearance: auto !important; accent-color: #D4AF37 !important; width: 18px !important; height: 18px !important; cursor: pointer !important; transform: scale(1.1) !important; }
.gradio-dropdown input, .gradio-dropdown select, .dropdown-menu, .options, input[type="text"], textarea { background-color: #0E0E12 !important; color: #FFF !important; border: 1px solid rgba(255,255,255,0.08) !important; border-radius: 12px !important; transition: border-color 0.3s !important;}
input:focus, textarea:focus, .gradio-dropdown:focus-within { border-color: #D4AF37 !important; outline: none !important; box-shadow: 0 0 10px rgba(212, 175, 55, 0.1) !important;}

/* SIDEBAR LUXURY */
.sidebar { background: #0E0E12 !important; border-right: 1px solid rgba(255,255,255,0.05) !important; padding: 25px 20px !important; }
.logo-container { text-align: center; margin-bottom: 30px; }
.logo-container img { margin-bottom: 15px; }

/* BOTÕES GERAIS */
button.primary { background: linear-gradient(135deg, #D4AF37, #AA7C11) !important; color: #000 !important; font-weight: bold !important; border-radius: 12px !important; border: none !important; transition: all 0.3s ease !important; }
button.primary:hover { transform: translateY(-2px); box-shadow: 0 6px 20px rgba(212, 175, 55, 0.3) !important; }
button.secondary { background: transparent !important; color: #AAA !important; border: 1px solid rgba(255,255,255,0.1) !important; border-radius: 12px !important; transition: all 0.3s !important; }
button.secondary:hover { border-color: #D4AF37 !important; color: #FFF !important; background: rgba(212, 175, 55, 0.05) !important;}

/* ABAS "PÍLULA" TECNOLÓGICAS */
.tabs { border: none !important; background: transparent !important; }
.tab-nav { background: rgba(255,255,255,0.02) !important; border: 1px solid rgba(255,255,255,0.05) !important; border-radius: 50px !important; padding: 6px !important; margin: 15px auto 25px auto !important; width: fit-content !important; display: flex !important; gap: 5px !important; }
.tab-nav button { background: transparent !important; color: #888 !important; border: none !important; border-radius: 50px !important; padding: 8px 20px !important; font-weight: 600 !important; font-size: 14px !important; transition: 0.3s !important;}
.tab-nav button.selected { background: #D4AF37 !important; color: #000 !important; box-shadow: 0 4px 15px rgba(212, 175, 55, 0.3) !important;}

/* CHAT MASTER (LIMPO E SEM BORDAS VAZANDO) */
.chatbot { background: transparent !important; border: none !important; }
.message-wrap { padding: 10px 0 !important; }
.message { font-size: 16px !important; line-height: 1.6 !important; padding: 18px 24px !important; max-width: 80% !important; border-radius: 20px !important; box-shadow: none !important;}
.message.user { background: #15151A !important; border: 1px solid rgba(212,175,55,0.2) !important; color: #FFF !important; margin-left: auto !important; border-bottom-right-radius: 4px !important; }
.message.bot { background: transparent !important; border: none !important; color: #E0E0E0 !important; margin-right: auto !important; padding-left: 0 !important; }

/* CAIXA DE DIGITAR PÍLULA FLUTUANTE */
.chat-container > div:last-child, .chat-container form { background: #0E0E12 !important; border: 1px solid rgba(255,255,255,0.1) !important; border-radius: 30px !important; padding: 6px 15px !important; margin: 0 auto 20px auto !important; max-width: 850px !important; box-shadow: 0 10px 30px rgba(0,0,0,0.5) !important; transition: border 0.3s !important;}
.chat-container form:focus-within { border-color: #D4AF37 !important; }

/* MICROFONE TECNOLÓGICO FIXO */
#btn-mic-titan { position: fixed !important; bottom: 25px !important; right: 25px !important; width: 65px !important; height: 65px !important; border-radius: 50% !important; background: rgba(14, 14, 18, 0.9) !important; backdrop-filter: blur(10px) !important; border: 1px solid rgba(212, 175, 55, 0.5) !important; box-shadow: 0 10px 30px rgba(0,0,0,0.8) !important; z-index: 999999 !important; display: flex !important; align-items: center !important; justify-content: center !important; cursor: pointer !important; transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275) !important;}
#btn-mic-titan svg { color: #D4AF37 !important; stroke: #D4AF37 !important; width: 28px !important; height: 28px !important; transition: all 0.3s !important; }
#btn-mic-titan:hover { transform: scale(1.1) !important; border-color: #D4AF37 !important; box-shadow: 0 15px 35px rgba(212,175,55,0.3) !important; background: #111 !important;}

@keyframes mic-record-anim { 
    0% { box-shadow: 0 0 0 0 rgba(255, 50, 50, 0.6); border-color: #ff3333; background: #220000; } 
    70% { box-shadow: 0 0 20px 20px rgba(255, 50, 50, 0); border-color: #ff5555; background: #330000; } 
    100% { box-shadow: 0 0 0 0 rgba(255, 50, 50, 0); border-color: #ff3333; background: #220000; } 
}
.mic-on { animation: mic-record-anim 1.5s infinite !important; }
.mic-on svg { stroke: #ff5555 !important; }

/* BARRAS DE ROLAGEM E CUSTOMIZAÇÕES GERAIS */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #333; border-radius: 10px; }
::-webkit-scrollbar-thumb:hover { background: #D4AF37; }

/* ================= MOBILE ================= */
@media screen and (max-width: 768px) {
    .tab-nav { width: 95% !important; overflow-x: auto !important; flex-wrap: nowrap !important; justify-content: flex-start !important; padding: 5px !important;}
    .tab-nav button { flex-shrink: 0 !important; font-size: 13px !important; padding: 6px 12px !important; }
    .message { max-width: 92% !important; font-size: 15px !important; padding: 14px 18px !important; }
    .chat-container > div:last-child, .chat-container form { width: 92% !important; border-radius: 24px !important; }
    #btn-mic-titan { bottom: 85px !important; right: 15px !important; width: 55px !important; height: 55px !important; }
    #btn-mic-titan svg { width: 24px !important; height: 24px !important; }
}
"""

JS_TITAN = """
function() {
    document.body.classList.add('dark');
    
    function initVoiceControl() {
        if (document.getElementById('btn-mic-titan')) return;
        
        const btn = document.createElement('button');
        btn.id = 'btn-mic-titan';
        btn.innerHTML = '<svg viewBox="0 0 24 24" fill="none" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2a3 3 0 0 0-3 3v7a3 3 0 0 0 6 0V5a3 3 0 0 0-3-3Z"></path><path d="M19 10v2a7 7 0 0 1-14 0v-2"></path><line x1="12" y1="19" x2="12" y2="22"></line></svg>';
        btn.title = "Comando de Voz";
        document.body.appendChild(btn);

        const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
        if(SpeechRecognition) {
            const recognition = new SpeechRecognition();
            recognition.lang = 'pt-BR';
            recognition.continuous = false;
            let isRecording = false;
            
            btn.onclick = (e) => {
                e.preventDefault();
                if(isRecording) recognition.stop();
                else recognition.start();
            };
            
            recognition.onstart = () => { 
                isRecording = true; 
                btn.classList.add('mic-on'); 
            };
            
            recognition.onresult = (event) => {
                let text = event.results[0][0].transcript;
                
                let activeInput = null;
                const textareas = document.querySelectorAll('textarea');
                textareas.forEach(ta => { if (ta.offsetParent !== null) activeInput = ta; });
                
                if (!activeInput) {
                    const inputs = document.querySelectorAll('input[type="text"]');
                    inputs.forEach(i => { if (i.offsetParent !== null) activeInput = i; });
                }

                if (activeInput) {
                    activeInput.value = activeInput.value ? activeInput.value + ' ' + text : text;
                    activeInput.dispatchEvent(new Event('input', { bubbles: true }));
                }
            };
            
            recognition.onend = () => { isRecording = false; btn.classList.remove('mic-on'); };
            recognition.onerror = () => { isRecording = false; btn.classList.remove('mic-on'); }
        } else {
            btn.style.display = 'none';
        }
    }
    
    setTimeout(initVoiceControl, 1000);
    setInterval(initVoiceControl, 2000);
}
"""

# ==========================================
# 7. CONSTRUÇÃO DA INTERFACE VISUAL (LIMPA E ESTRUTURADA)
# ==========================================
with gr.Blocks(title="Código de Ouro", theme=tema_titan, css=CSS_TITAN, fill_height=True) as interface:
    id_sessao_atual = gr.State(f"Chat_{datetime.now().strftime('%d%m_%H%M%S')}")

    with gr.Row():
        
        # BARRA LATERAL (Agora guarda as configurações para despoluir a tela)
        with gr.Column(scale=2, min_width=280, elem_classes="sidebar"):
            gr.HTML(f"""
            <div class="logo-container">
                {TAG_LOGO}
            </div>
            """)
            btn_novo = gr.Button("➕ Novo Atendimento", variant="primary")
            
            gr.Markdown("<br>### ⚙️ Engine da IA")
            persona = gr.Dropdown(choices=["Assistente Padrão", "Gênio do Marketing", "Analista de Dados", "Estrategista de Negócios"], value="Assistente Padrão", label="Especialidade", interactive=True)
            net = gr.Checkbox(label="🌐 Pesquisa Web Integrada", value=False)
            btn_exportar = gr.Button("💾 Exportar Documento", variant="secondary")
            
            gr.Markdown("<br>### 📜 Histórico Criptografado")
            lista_chats = gr.Dropdown(choices=listar_sessoes_chat(), label="Sessões Salvas", interactive=True)
            with gr.Row():
                btn_load = gr.Button("Abrir", variant="secondary")
                btn_atualizar = gr.Button("Atualizar", variant="secondary")
            btn_atualizar.click(lambda: gr.update(choices=listar_sessoes_chat()), None, lista_chats)

        # ÁREA CENTRAL DE ALTA PERFORMANCE (Apenas o conteúdo importa)
        with gr.Column(scale=8):
            with gr.Tabs():
                
                with gr.TabItem("💬 Console Central"):
                    # O ChatInterface operando sozinho, sem menus em cima, no melhor estilo ChatGPT
                    chat = gr.ChatInterface(
                        fn=responder_chat_central, multimodal=True, additional_inputs=[persona, net, id_sessao_atual],
                        chatbot=gr.Chatbot(show_label=False), textbox=gr.MultimodalTextbox(placeholder="Descreva seu projeto, anexe documentos ou imagens...", container=False)
                    )
                    arq_exportado = gr.File(label="Documento Gerado", visible=False)
                    
                    # Conexões dos botões da Sidebar com a tela de Chat
                    btn_exportar.click(exportar_conversa_docx, chat.chatbot, arq_exportado).then(lambda: gr.update(visible=True), None, arq_exportado)
                    btn_load.click(carregar_sessao_chat, lista_chats, [chat.chatbot, id_sessao_atual])
                    btn_novo.click(iniciar_novo_chat, None, [chat.chatbot, id_sessao_atual, lista_chats])

                with gr.TabItem("📑 Analisador de Data Room"):
                    gr.Markdown("### 🧠 Extração e Análise Profunda")
                    with gr.Row():
                        with gr.Column(scale=4):
                            files = gr.File(label="Upload de PDFs, Excel ou Word", file_count="multiple")
                            inst = gr.Textbox(label="Diretriz da Análise", placeholder="O que eu devo buscar ou estruturar com base nestes arquivos?", lines=4)
                            btn_doc = gr.Button("Iniciar Processamento", variant="primary")
                        with gr.Column(scale=6):
                            res_doc = gr.Textbox(label="Relatório Analítico", lines=15)
                    btn_doc.click(gerar_dossie_lote, [files, inst], [res_doc])

                with gr.TabItem("🗂️ Cofre de Ativos"):
                    btn_att_gal = gr.Button("🔄 Sincronizar Galeria", variant="primary")
                    gal = gr.Gallery(columns=4, height="auto", label="Mídias Recentes")
                    btn_att_gal.click(atualizar_galeria_imagens, None, gal)

# ==========================================
# 8. LANÇAMENTO SEGURO
# ==========================================
usuarios = []
for i in ["", "_1", "_2", "_3", "_4", "_5", "_6", "_7", "_8", "_9"]:
    u = os.environ.get(f"LOGIN_USUARIO{i}") or os.environ.get(f"USUARIO{i}")
    s = os.environ.get(f"LOGIN_SENHA{i}") or os.environ.get(f"SENHA{i}")
    if u and s: usuarios.append((u, s))

launch_args = {
    "server_name": "0.0.0.0", 
    "server_port": int(os.environ.get("PORT", 10000)),
    "auth": usuarios if usuarios else None, 
    "auth_message": LOGIN_HACK,
    "js": JS_TITAN, 
    "head": PWA_HEAD
}

if caminho_logo: launch_args["favicon_path"] = caminho_logo

interface.launch(**launch_args)
